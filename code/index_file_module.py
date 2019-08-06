# -*- coding: utf-8 -*-


from os import listdir
import jieba
import sqlite3
import configparser
import json
import docx


class Doc:
    docid = 0
    tf = 0
    ld = 0

    def __init__(self, docid, tf, ld):
        self.docid = docid
        self.tf = tf  # 该词项出现的次数，即词项频率(tf)
        self.ld = ld  # 该文档的长度(ld)

    def __repr__(self):
        return (str(self.docid) + '\t' + str(self.tf) + '\t' + str(self.ld))

    def __str__(self):
        return (str(self.docid) + '\t' + str(self.tf) + '\t' + str(self.ld))


class IndexModule:
    stop_words = set()
    postings_lists = {}

    config_path = ''
    config_encoding = ''

    def __init__(self, config_path, config_encoding):
        self.config_path = config_path
        self.config_encoding = config_encoding
        config = configparser.ConfigParser()
        config.read(config_path, config_encoding)
        f = open(config['DEFAULT']['stop_words_path'], encoding=config['DEFAULT']['stop_words_encoding'])
        words = f.read()
        self.stop_words = set(words.split('\n'))

    def is_number(self, s):
        try:
            float(s)
            return True
        except ValueError:
            return False

    def clean_list(self, seg_list):
        '''

        :param seg_list:
        :return:
        '''
        cleaned_dict = {}
        n = 0
        for i in seg_list:
            i = i.strip().lower()
            if i != '' and not self.is_number(i) and i not in self.stop_words:
                n = n + 1
                if i in cleaned_dict:
                    cleaned_dict[i] = cleaned_dict[i] + 1
                else:
                    cleaned_dict[i] = 1
        return n, cleaned_dict

    def write_files_postings_to_db(self, db_path):
        '''
        写入数据库
        :param db_path:
        :return:
        '''
        conn = sqlite3.connect(db_path)
        c = conn.cursor()

        c.execute('''DROP TABLE IF EXISTS files_postings''')
        c.execute('''CREATE TABLE files_postings
                     (term TEXT PRIMARY KEY, df INTEGER, docs TEXT)''')

        for key, value in self.postings_lists.items():
            doc_list = '\n'.join(map(str, value[1]))
            t = (key, value[0], doc_list)
            c.execute("INSERT INTO files_postings VALUES (?, ?, ?)", t)

        conn.commit()
        conn.close()

    def construct_files_postings_lists(self):
        '''

        :return:
        '''
        config = configparser.ConfigParser()
        config.read(self.config_path, self.config_encoding)
        files = listdir(config['DEFAULT']['doc_dir_path'])  # 所有的网页文件
        id = 1175  # 紧随news的id
        AVG_L = 0
        for item in files:
            # TODO
            print(item)
            doc = docx.Document(config['DEFAULT']['doc_dir_path'] + item)
            fullText = []
            for p in doc.paragraphs:  # 迭代docx文档里面的每一个段落
                fullText.append(p.text)  # 保存每一个段落的文本
            t = (id, item, '\n'.join(fullText))
            title = item
            body = '\n'.join(fullText)
            seg_list = jieba.lcut(title + '。' + body, cut_all=False)  # 分词，标题+正文
            ld, cleaned_dict = self.clean_list(seg_list)
            AVG_L = AVG_L + ld
            for key, value in cleaned_dict.items():
                d = Doc(id, value, ld)
                if key in self.postings_lists:  # 词项在不同文档中出现的次数，即文档频率(df)
                    self.postings_lists[key][0] = self.postings_lists[key][0] + 1  # df++
                    self.postings_lists[key][1].append(d)
                else:
                    self.postings_lists[key] = [1, [d]]  # [df, [Doc]]

            id += 1

        AVG_L = AVG_L / id
        config.set('DEFAULT', 'N', str(id))
        config.set('DEFAULT', 'avg_l', str(AVG_L))
        with open(self.config_path, 'w', encoding=self.config_encoding) as configfile:
            config.write(configfile)
        self.write_files_postings_to_db(config['DEFAULT']['db_path'])

    def construct_files_lists(self, db_path):
        '''
        附件 数据库表格
        :param db_path:
        :return:
        '''
        config = configparser.ConfigParser()
        config.read(self.config_path, self.config_encoding)
        files = listdir(config['DEFAULT']['doc_dir_path'])  # 所有的网页文件

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        id = 1175
        rank = 1
        for item in files:
            # TODO
            print(item)
            doc = docx.Document(config['DEFAULT']['doc_dir_path'] + item)
            fullText = []
            for p in doc.paragraphs:  # 迭代docx文档里面的每一个段落
                fullText.append(p.text)  # 保存每一个段落的文本
            if id - 1175 < 80:
                rank = 1
            elif id - 1175 < 160:
                rank = 2
            elif id - 1175 < 240:
                rank = 3
            else:
                rank = 4
            t = (id, item, '\n'.join(fullText), '', '', rank)
            c.execute("INSERT INTO news VALUES (?, ?, ?,?,?,?)", t)
            id += 1
        conn.commit()
        conn.close()


if __name__ == "__main__":
    im = IndexModule('../config.ini', 'utf-8')
    # im.construct_files_postings_lists()
    im.construct_files_lists('D:\\Work\\IR\\Lab3_search_engine\\data\\ir.db')
